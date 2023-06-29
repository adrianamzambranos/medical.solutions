from docx import Document
import speech_recognition as sr
from pydub import AudioSegment

def guardar_texto_en_word(datos, nombre_archivo):
    # Crear un nuevo documento de Word
    document = Document()

    # Agregar los encabezados de sección y los párrafos correspondientes
    document.add_heading('Nombre del paciente', level=1)
    document.add_paragraph(datos['nombre_paciente'])

    document.add_heading('Nombre del médico', level=1)
    document.add_paragraph(datos['nombre_medico'])

    document.add_heading('Fecha', level=1)
    document.add_paragraph(datos['fecha'])

    document.add_heading('Hora', level=1)
    document.add_paragraph(datos['hora'])

    document.add_heading('Datos del paciente', level=1)
    document.add_paragraph(datos['datos_paciente'])

    document.add_heading('Diagnóstico', level=1)
    document.add_paragraph(datos['diagnostico'])

    document.add_heading('Receta', level=1)
    document.add_paragraph(datos['receta'])

    document.add_heading('Recomendaciones generales', level=1)
    document.add_paragraph(datos['recomendaciones_generales'])

    # Guardar el documento como archivo de Word
    document.save(nombre_archivo)


# Ruta del archivo MP3
mp3_file = "Paciente.mp3"

# Ruta de salida para el archivo WAV
wav_file = "salida.wav"

# Convertir MP3 a WAV
audio = AudioSegment.from_mp3(mp3_file)
audio.export(wav_file, format="wav")

# Crear un objeto de reconocimiento de voz
r = sr.Recognizer()

# Transcribir el archivo WAV
with sr.AudioFile(wav_file) as source:
    # Leer el archivo de audio
    audio_data = r.record(source)
    # Realizar la transcripción utilizando el servicio de reconocimiento de voz
    text = r.recognize_google(audio_data, language="es-MX")

# Imprimir el texto transcrito
print(text)

# Procesar el texto transcrito y clasificarlo en diferentes secciones
datos = {
    'nombre_paciente': 'John Doe',
    'nombre_medico': 'Dr. Smith',
    'fecha': '29 de junio de 2023',
    'hora': '10:00 AM',
    'datos_paciente': 'Lorem ipsum dolor sit amet, consectetur adipiscing elit.',
    'diagnostico': 'Lorem ipsum dolor sit amet, consectetur adipiscing elit.',
    'receta': 'Lorem ipsum dolor sit amet, consectetur adipiscing elit.',
    'recomendaciones_generales': 'Lorem ipsum dolor sit amet, consectetur adipiscing elit.'
}

# Guardar el texto clasificado en un archivo de Word
nombre_archivo = "NewPaciente.docx"
guardar_texto_en_word(datos, nombre_archivo)
